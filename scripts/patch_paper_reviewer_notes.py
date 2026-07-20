"""Reviewer round: class-imbalance note, temporal-leakage justification, language polish."""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "AWID3_Paper_IEEE_Access.docx"
BACKUP = ROOT / "AWID3_Paper_IEEE_Access_pre_reviewer_notes.docx"


def set_text(p, text: str) -> None:
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def replace_in_para(doc, needle: str, old: str, new: str) -> None:
    for p in doc.paragraphs:
        if needle in p.text and old in p.text:
            set_text(p, p.text.replace(old, new))
            return
    raise RuntimeError(f"not found: {needle!r} / {old!r}")


def append_to_para(doc, needle: str, addition: str) -> None:
    for p in doc.paragraphs:
        if needle in p.text:
            set_text(p, p.text.rstrip() + " " + addition)
            return
    raise RuntimeError(f"not found: {needle!r}")


def main() -> None:
    if not BACKUP.exists():
        shutil.copy2(DOCX, BACKUP)
    doc = Document(str(DOCX))

    # 1) Class-imbalance handling (Section IV-C)
    append_to_para(
        doc,
        "Re-running it reproduces the dataset exactly, which is the point.",
        "We deliberately avoid synthetic oversampling (e.g., SMOTE): capping the majority Normal class at "
        "20,000 already narrows the imbalance, stratified folds preserve each class's proportion so even the "
        "smallest categories (SQL_Injection, RogueAP, (Re)Assoc) appear in every train and test partition, and "
        "we report macro-averaged F1 so minority classes weigh equally in the headline metric. Tree ensembles are "
        "comparatively robust to residual imbalance, and LightGBM additionally uses class-balanced weighting; no "
        "per-class resampling or focal loss is applied.",
    )

    # 2) Temporal-leakage justification (Section V)
    append_to_para(
        doc,
        "All runs are on a commodity CPU workstation with no GPU, so the deep models use CPU builds.",
        "We use random stratified folds rather than TimeSeriesSplit or GroupKFold on purpose: each feature vector "
        "is computed independently from a single non-overlapping (tumbling) frame window and carries no cross-window "
        "temporal state, so there is no look-ahead dependency between rows for a random split to exploit. The "
        "capture-environment leakage that does matter is handled separately by the same-capture curation of "
        "Section IV-A and quantified against the naive across-capture split.",
    )

    # 3) Language polish
    replace_in_para(
        doc,
        "Wi-Fi intrusion detectors are routinely reported",
        "Those numbers deserve a second look.",
        "However, these reported figures warrant a rigorous re-examination.",
    )
    replace_in_para(
        doc,
        "This paper takes the opposite stance.",
        "We do not chase the highest score.",
        "Rather than merely maximizing a headline score, this study prioritizes methodological rigor.",
    )
    replace_in_para(
        doc,
        "Figures 9 and 10 project the same thirty-four clean features",
        "Numbers can be argued with; a picture is harder to dismiss.",
        "While numerical metrics provide quantitative evidence, a visual projection offers clear qualitative "
        "confirmation of capture-environment leakage.",
    )

    doc.save(str(DOCX))
    print("Applied: imbalance note, temporal-leakage justification, 3 language edits.")


if __name__ == "__main__":
    main()
