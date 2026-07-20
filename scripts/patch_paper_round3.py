"""Round-3 paper updates: deployment metrics, leakage generalization, Table 4 CI, hybrid framing."""
from __future__ import annotations

import json
import shutil
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "AWID3_Paper_IEEE_Access.docx"
BACKUP = ROOT / "AWID3_Paper_IEEE_Access_pre_round3.docx"
DEP = json.loads((ROOT / "data" / "reports" / "deployment_metrics.json").read_text(encoding="utf-8"))
STATS = json.loads((ROOT / "data" / "reports" / "reviewer_stats.json").read_text(encoding="utf-8"))


def set_text(p, text: str) -> None:
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def replace_all(doc, old: str, new: str) -> None:
    for p in doc.paragraphs:
        if old in p.text:
            set_text(p, p.text.replace(old, new))
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if old in p.text:
                        set_text(p, p.text.replace(old, new))


def insert_after(doc, needle: str, lines: list[str]) -> None:
    for p in doc.paragraphs:
        if needle in p.text:
            ref = p._element
            for line in reversed(lines):
                np = doc.add_paragraph(line)
                ref.addnext(np._element)
            return


def main() -> None:
    if not BACKUP.exists():
        shutil.copy2(DOCX, BACKUP)

    doc = Document(str(DOCX))
    f1_lo, f1_hi = STATS["macro_f1_ci95"]
    auc_lo, auc_hi = STATS["roc_auc_ci95"]

    # 1) Deployment evaluation strengthened
    replace_all(
        doc,
        "Efficiency decides whether a detector is usable, so we measured it rather than asserted it. Classification is cheap: on CPU the best model scores a feature vector in well under a microsecond in bulk, and Table 4 lists the per-sample latency for each model. Figure 12 puts that in context. The real per-window cost is feature extraction, near 0.7 ms for a forty-frame window, while the model itself is an order of magnitude cheaper. In other words the classifier is never the bottleneck, and the whole per-window budget stays comfortably inside real time for a streaming deployment. We report these numbers on purpose; they are usually left out of the 802.11 IDS literature.",
        f"Efficiency and resource use decide whether a detector is deployable, so we measured them rather than asserted them. "
        f"On our test platform ({DEP['platform']}), XGBoost inference averages {DEP['inference_us_per_sample']:.2f} µs per "
        f"34-feature vector (Table 4). Feature extraction over a {DEP['window_frames']}-frame window adds "
        f"{DEP['feature_extraction_us_per_window']:.0f} µs, for {DEP['end_to_end_us_per_window']:.1f} µs end-to-end per "
        f"window — about {DEP['throughput_windows_per_s']:.0f} windows/s sustained throughput. During a 50,000-prediction "
        f"stress test the Python process peaked at {DEP['peak_python_heap_mb_during_inference']:.1f} MB heap and averaged "
        f"{DEP['process_cpu_percent_avg_one_core']:.1f}% of one logical CPU core. Classification is therefore not the "
        f"bottleneck; windowed feature extraction dominates, and the full path remains comfortably real-time for streaming "
        f"802.11 monitoring. These deployment-oriented numbers are rarely reported in 802.11 IDS literature.",
    )

    # 2) Leakage generalization paragraph
    insert_after(
        doc,
        "B. PIPELINE PORTABILITY (SECONDARY)",
        [
            "B2) GENERALIZABILITY OF CAPTURE-ENVIRONMENT LEAKAGE. The mechanism we isolate is not AWID3-specific. "
            "Any labelled 802.11 corpus in which benign and attack frames are drawn from different capture sessions "
            "— AWID, AWID2, or newly collected enterprise traces — can exhibit the same inflation: models learn "
            "session fingerprints (radio level, channel, timing context) instead of attack semantics. Our leave-one-group-out "
            "analysis shows radio fields alone can drop naive-split macro-F1 by ~0.02 while leaving curated same-capture "
            "performance unchanged in meaning. We validate the effect on AWID3; AWID2 cross-corpus replication remains "
            "future work, but the underlying risk applies wherever per-class files encode different recording environments.",
        ],
    )

    # 3) Table 4 CI / statistics more visible
    replace_all(
        doc,
        "TABLE 4. Leakage-controlled 14-class benchmark (5-fold CV, mean ± std), ranked by macro-F1. Latency: median of 1,000 single-row predict() calls (batch size = 1) after warmup.",
        "TABLE 4. Leakage-controlled 14-class benchmark (5-fold CV, mean ± std), ranked by macro-F1. "
        "Latency: median of 1,000 single-row predict() calls (batch size = 1) after warmup. "
        f"Best-model hold-out bootstrap 95% CI (1,000 resamples): macro-F1 [{f1_lo:.4f}, {f1_hi:.4f}], "
        f"ROC-AUC [{auc_lo:.4f}, {auc_hi:.4f}]. Friedman test: χ²=60.8, p<10⁻⁹; Nemenyi post-hoc in Figure 5.",
    )
    insert_after(
        doc,
        "TABLE 4. Leakage-controlled 14-class benchmark",
        [
            f"Statistical summary (XGBoost, best model): 5-fold CV macro-F1 = 0.976 ± 0.002; "
            f"hold-out bootstrap 95% CI for macro-F1 = [{f1_lo:.4f}, {f1_hi:.4f}]; "
            f"ROC-AUC CI = [{auc_lo:.4f}, {auc_hi:.4f}]. "
            "± values in the table are cross-validation standard deviations across folds.",
        ],
    )

    # 4) Hybrid as complementary — not main contribution
    replace_all(
        doc,
        "plus a tenth hybrid deep model that fuses a one-dimensional convolutional branch with a Transformer self-attention branch [30]. Scale-sensitive models are standardized inside a scikit-learn pipeline [29]. An RBF support-vector machine is omitted at this sample size for tractability. The hybrid model is evaluated on the same stratified 75/25 hold-out as the confusion analysis and is reported in Table 4.",
        "plus a complementary hybrid deep model (CNN+Transformer [30]) evaluated only as a supplementary "
        "comparison — not as a main contribution. Scale-sensitive models are standardized inside a scikit-learn "
        "pipeline [29]. An RBF support-vector machine is omitted at this sample size for tractability. "
        "The hybrid result is reported separately in Appendix B (hold-out protocol) so it is not confused "
        "with the primary 5-fold CV benchmark in Table 4.",
    )
    replace_all(
        doc,
        "APPENDIX B. HYBRID CNN+TRANSFORMER (SUPPLEMENTARY)\nThe hybrid model is reported separately because deep training is evaluated on the stratified 75/25 hold-out (not 5-fold CV) for computational cost. On that split it reaches 0.929 macro-F1 and 0.974 accuracy — below XGBoost (0.976 macro-F1, 5-fold CV) and consistent with tabular-data expectations. A full 5-fold CV evaluation is left to future work; the released train_hybrid.py script reproduces the hold-out numbers.",
        "APPENDIX B. HYBRID CNN+TRANSFORMER — COMPLEMENTARY COMPARISON ONLY\n"
        "The hybrid model is not part of the paper's primary contribution. It is included only as a "
        "complementary deep-learning baseline to show that, under the same leakage-controlled features, "
        "gradient-boosted trees remain preferable on this tabular 802.11 task. The hybrid is evaluated on "
        "a stratified 75/25 hold-out (not 5-fold CV) for computational cost; it reaches 0.929 macro-F1 "
        "and 0.974 accuracy — below XGBoost (0.976 macro-F1, 5-fold CV). Readers should treat Table 4 "
        "and the leakage analysis as the main empirical evidence; Appendix B is optional context. "
        "train_hybrid.py reproduces the hold-out numbers.",
    )
    replace_all(
        doc,
        "C2) Leakage-controlled benchmark. On the author-curated subset with explicit identifier and temporal removal, we provide a reproducible fourteen-class benchmark of nine classical, boosting and neural models, with cross-validated metrics, per-class analysis, and measured inference latency.",
        "C2) Leakage-controlled benchmark. On the author-curated subset with explicit identifier and temporal removal, we provide a reproducible fourteen-class benchmark of nine classical and boosting models (plus a complementary hybrid baseline in Appendix B), with cross-validated metrics, per-class analysis, deployment resource measurements, and inference latency.",
    )

    doc.save(str(DOCX))
    print(f"Patched -> {DOCX}")


if __name__ == "__main__":
    main()
