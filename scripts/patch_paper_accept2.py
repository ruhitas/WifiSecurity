# -*- coding: utf-8 -*-
"""Acceptance round 2: positioning sentence, full-archive rationale, WPA4 specificity, versions, de-duplication."""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "AWID3_Paper_IEEE_Access.docx"
BACKUP = ROOT / "AWID3_Paper_IEEE_Access_pre_accept2.docx"


def set_text(p, text: str) -> None:
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def edit(doc, anchor: str, old: str, new: str) -> None:
    for p in doc.paragraphs:
        if anchor in p.text and old in p.text:
            set_text(p, p.text.replace(old, new, 1))
            return
    raise RuntimeError(f"not found: {anchor!r} / {old!r}")


def edit_cell(doc, old: str, new: str) -> None:
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if old in p.text:
                        set_text(p, p.text.replace(old, new, 1))
                        return
    raise RuntimeError(f"cell not found: {old!r}")


def main() -> None:
    if not BACKUP.exists():
        shutil.copy2(DOCX, BACKUP)
    doc = Document(str(DOCX))

    # 1) VII.A positioning vs recent SOTA
    edit(
        doc,
        "A. NOVELTY AND POSITIONING" if False else "Two things set the work apart.",
        "both matter to anyone who has to defend a deployed detector.",
        "both matter to anyone who has to defend a deployed detector. While recent temporal-context and hybrid "
        "models report higher headline numbers on coarser class sets or without a measured leakage split, it is "
        "precisely our controlled protocol that exposes the true fourteen-class difficulty; a like-for-like "
        "ranking would require re-running those models under the same leakage control, which we leave as an open "
        "invitation to the community.",
    )

    # 2) Full-archive rationale (computational / memory cost)
    edit(
        doc,
        "Full-archive protocol.",
        "and is a priority extension alongside cross-wireless-dataset validation.",
        "and is a priority extension alongside cross-wireless-dataset validation. It is deferred here for a "
        "practical reason: same-capture sampling across the full 43 GB archive is memory- and compute-intensive, "
        "whereas the curated subset already isolates the leakage effect cleanly.",
    )

    # 3) WPA4 specificity (hedged)
    edit(
        doc,
        "The third is the moving target of the standard.",
        "future amendments to the standard will follow",
        "future amendments to the standard \u2014 a prospective WPA4 and beyond \u2014 will follow",
    )

    # 4) Software versions completeness (SHAP, PyTorch)
    edit(
        doc,
        "The workstation is an Intel Core Ultra 7 155H",
        "scikit-learn 1.9, XGBoost 3.2, LightGBM 4.6, NumPy 2.2 and pandas.",
        "scikit-learn 1.9, XGBoost 3.2, LightGBM 4.6, SHAP 0.51, PyTorch 2.13 (CPU), NumPy 2.2 and pandas.",
    )

    # 5) Reduce "leakage-controlled" repetition (non-anchor prose/captions)
    edit(
        doc,
        "We studied machine-learning intrusion detection",
        "a rigorous, leakage-controlled picture emerges:",
        "a rigorous, controlled picture emerges:",
    )
    edit(
        doc,
        "FIGURE 6. Correlation matrix",
        "Correlation matrix of the 34 leakage-controlled features",
        "Correlation matrix of the 34 curated features",
    )
    edit(
        doc,
        "Evaluation scope.",
        "its curated subset, though leakage-controlled, is one testbed",
        "its curated subset, though controlled for leakage, is one testbed",
    )

    doc.save(str(DOCX))
    print("Applied acceptance round 2 edits.")


if __name__ == "__main__":
    main()
