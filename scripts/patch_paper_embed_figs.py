"""Embed the missing Figures 13-16 above their existing captions."""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "AWID3_Paper_IEEE_Access.docx"
BACKUP = ROOT / "AWID3_Paper_IEEE_Access_pre_embedfigs.docx"
FIGDIR = ROOT / "data" / "reports" / "figures"

# caption prefix -> image file
MAP = {
    "FIGURE 13.": FIGDIR / "fig_leakage_group_ablation.png",
    "FIGURE 14.": FIGDIR / "fig_shap_local_Normal.png",
    "FIGURE 15.": FIGDIR / "fig_shap_local_Krack.png",
    "FIGURE 16.": FIGDIR / "fig_shap_local_Deauth.png",
}


def main() -> None:
    if not BACKUP.exists():
        shutil.copy2(DOCX, BACKUP)

    doc = Document(str(DOCX))
    done = []
    for prefix, img in MAP.items():
        if not img.exists():
            raise FileNotFoundError(img)
        cap = next(p for p in doc.paragraphs if p.text.strip().startswith(prefix))
        pic_p = cap.insert_paragraph_before()
        pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic_p.add_run().add_picture(str(img), width=Inches(3.12))
        done.append(prefix)

    doc.save(str(DOCX))
    print("Embedded:", ", ".join(done))


if __name__ == "__main__":
    main()
