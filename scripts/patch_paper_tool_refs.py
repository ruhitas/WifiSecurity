# -*- coding: utf-8 -*-
"""Restore well-known tool names in the deployment algorithm and add cited web references."""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "AWID3_Paper_IEEE_Access.docx"
BACKUP = ROOT / "AWID3_Paper_IEEE_Access_pre_toolrefs.docx"

LQ, RQ = "\u201c", "\u201d"

# temporary numbers; a following renumber pass fixes ordering
NEW_REFS = [
    f"[43] Redis Ltd., {LQ}Redis: The real-time data platform,{RQ} 2026. [Online]. "
    "Available: https://redis.io. Accessed: Jul. 1, 2026.",
    f"[44] Apache Software Foundation, {LQ}Apache Kafka: A distributed event streaming platform,{RQ} 2026. "
    "[Online]. Available: https://kafka.apache.org. Accessed: Jul. 2, 2026.",
    f"[45] Microsoft, {LQ}Microsoft SQL Server,{RQ} 2026. [Online]. "
    "Available: https://www.microsoft.com/sql-server. Accessed: Jul. 3, 2026.",
    f"[46] Elastic N.V., {LQ}Elasticsearch: The heart of the Elastic Stack,{RQ} 2026. [Online]. "
    "Available: https://www.elastic.co/elasticsearch. Accessed: Jul. 4, 2026.",
]


def set_text(p, text: str) -> None:
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def replace_cell(doc, old: str, new: str) -> None:
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if old in p.text:
                        set_text(p, p.text.replace(old, new, 1))
                        return
    raise RuntimeError(f"cell not found: {old!r}")


def main() -> None:
    if not BACKUP.exists():
        shutil.copy2(DOCX, BACKUP)
    doc = Document(str(DOCX))

    # restore brand names + attach citations
    replace_cell(doc, "cache v; publish v to the streaming bus",
                 "cache v in Redis [43]; publish v to Kafka [44]")
    replace_cell(doc, "persist the event; index the alert for search",
                 "persist event to MSSQL [45]; index alert in Elasticsearch [46]")

    # append the four web references after the last reference entry
    import re
    last_ref = None
    for p in doc.paragraphs:
        if re.match(r"^\[\d+\]", p.text.strip()):
            last_ref = p
    if last_ref is None:
        raise RuntimeError("no reference entries found")
    el = last_ref._element
    for line in reversed(NEW_REFS):
        np = doc.add_paragraph(line)
        el.addnext(np._element)

    doc.save(str(DOCX))
    print("Restored tool names with citations [43]-[46]; appended web references.")


if __name__ == "__main__":
    main()
