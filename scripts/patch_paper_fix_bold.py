# -*- coding: utf-8 -*-
"""Restore standard formatting: bold lead-in label + normal body for paragraphs
that were accidentally made fully bold by earlier text edits."""
from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "AWID3_Paper_IEEE_Access.docx"
BACKUP = ROOT / "AWID3_Paper_IEEE_Access_pre_boldfix.docx"

CAPTION_RE = re.compile(r"^((?:FIGURE|TABLE) (?:\d+|A1)\.\s*)")
LEADIN_RE = re.compile(
    r"^(C[1-5]\) [^.]+\.\s*"
    r"|Evaluation scope\.\s*"
    r"|Class imbalance\.\s*"
    r"|Hardware\.\s*"
    r"|Full-archive protocol\.\s*)"
)


def rebuild(doc, p, bold_part: str, rest: str, style_name: str | None = None) -> None:
    text_runs = [r for r in p.runs]
    if not text_runs:
        return
    first = text_runs[0]
    first.text = bold_part
    first.bold = True
    for r in text_runs[1:]:
        r.text = ""
    if rest:
        nr = p.add_run(rest)
        nr.bold = False
    if style_name:
        p.style = doc.styles[style_name]


def all_bold(p) -> bool:
    runs = [r for r in p.runs if r.text.strip()]
    return bool(runs) and all(r.bold for r in runs)


def main() -> None:
    if not BACKUP.exists():
        shutil.copy2(DOCX, BACKUP)
    doc = Document(str(DOCX))

    fixed = []
    for i, p in enumerate(doc.paragraphs):
        t = p.text
        if not t.strip():
            continue

        m = CAPTION_RE.match(t.strip())
        if m:
            is_wrong_style = p.style.name == "Normal"
            if all_bold(p) or is_wrong_style:
                label = m.group(1)
                rest = t.strip()[len(label):]
                target = None
                if is_wrong_style:
                    target = "Fig Caption" if label.startswith("FIGURE") else "Table Caption"
                rebuild(doc, p, label, rest, target)
                fixed.append((i, label.strip()))
            continue

        m = LEADIN_RE.match(t.strip())
        if m and all_bold(p):
            label = m.group(1)
            rest = t.strip()[len(label):]
            rebuild(doc, p, label, rest)
            fixed.append((i, label.strip()[:40]))

    doc.save(str(DOCX))
    for i, lbl in fixed:
        print(f"fixed P{i}: {lbl}")
    print(f"total fixed: {len(fixed)}")


if __name__ == "__main__":
    main()
