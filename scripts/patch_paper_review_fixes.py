"""Apply reviewer major-revision fixes to AWID3_Paper_IEEE_Access.docx."""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "AWID3_Paper_IEEE_Access.docx"
BACKUP = ROOT / "AWID3_Paper_IEEE_Access_pre_review_fix.docx"
FEATURES = (ROOT / "config" / "leakage_controlled_features.txt").read_text(encoding="utf-8").strip().splitlines()
REPO_URL = "https://github.com/ruhitas/awid3-leakage-aware-ids"

REPLACEMENTS: list[tuple[str, str]] = [
    ("forty features", "thirty-four features"),
    ("Forty features", "Thirty-four features"),
    ("the forty features", "the thirty-four features"),
    ("The forty features", "The thirty-four features"),
    ("40 leakage-controlled features", "34 leakage-controlled features"),
    ("40 features", "34 features"),
    ("40 clean features", "34 clean features"),
    ("40 behavioural features", "34 behavioural features"),
    ("40-feature set", "34-feature set"),
    ("40-feature", "34-feature"),
    ("cleaned 40-feature set", "cleaned 34-feature set"),
    (
        "We drop any feature that names the capture or the endpoints instead of describing the attack. "
        "That means every timestamp, time-delta, TSF and MAC-time field, the frame numbers, and the MAC "
        "and IP address columns. Transport ports stay, since a port is behaviour, not identity. Constant "
        "columns go too. Forty features remain. Four transport time-delta and time-relative columns survive "
        "in some earlier pipelines; we remove them here, and doing so barely moves the score (Section VI-C), "
        "which tells us the curated subset was already clean of the worse offenders.",
        "We drop any feature that names the capture or the endpoints instead of describing the attack. "
        "That means every timestamp, time-delta, TSF and MAC-time field, the frame numbers, MAC and IP "
        "address columns, and per-frame or transport sequence counters (wlan.seq, tcp.seq, tcp.ack and their "
        "raw variants). Transport ports stay, since a port is behaviour, not identity. Constant columns go "
        "too. Thirty-four behavioural features remain (listed in the reproducibility package and Table 2 note). "
        "Removing sequence fields does not materially change the benchmark (Section VI-C).",
    ),
    ("which leaves forty features", "which leaves thirty-four features"),
    ("keep the 40 behavioural features", "keep the 34 behavioural features"),
    (
        "C3) Explainability. We attach SHAP reasoning reports to detections and check that the most "
        "influential features agree with domain knowledge — UDP ports for SSDP, channel and sequence for KRACK.",
        "C3) Explainability. We attach SHAP reasoning reports to detections and check that the most "
        "influential features agree with domain knowledge — UDP ports for SSDP, channel and frame-control "
        "fields for KRACK.",
    ),
    (
        "KRACK is driven by channel, frequency and sequence number.",
        "KRACK is driven by channel, frequency and frame-control subtype.",
    ),
    (
        "We also bootstrap the best model's macro ROC-AUC over a thousand resamples; the 95% interval is "
        "tight, [0.9996, 0.9997]. The reading is plain.",
        "We also bootstrap macro ROC-AUC on held-out predictions (1,000 resamples); the point estimate is "
        "0.9997 with a 95% interval of [0.9996, 0.9997], consistent with the five-fold CV mean of 0.9998 "
        "in Table 4. The reading is plain.",
    ),
    (
        "the macro-average AUC is 0.9998.",
        "the macro-average AUC is 0.9998 (CV); held-out bootstrap mean 0.9997.",
    ),
    (
        "The supervised pool is nine models: logistic regression, k-nearest neighbours, a decision tree, "
        "random forest, extremely randomized trees, gradient boosting, a multilayer perceptron, and the "
        "boosting libraries XGBoost [27] and LightGBM [28]. Scale-sensitive models are standardized inside "
        "a scikit-learn pipeline [29]. An RBF support-vector machine is omitted at this sample size for "
        "tractability. We also train a hybrid deep model that fuses a one-dimensional convolutional branch "
        "with a Transformer self-attention branch [30].",
        "The supervised pool is nine classical and boosting models — logistic regression, k-nearest "
        "neighbours, a decision tree, random forest, extremely randomized trees, gradient boosting, a "
        "multilayer perceptron, XGBoost [27] and LightGBM [28] — plus a tenth hybrid deep model that fuses "
        "a one-dimensional convolutional branch with a Transformer self-attention branch [30]. "
        "Scale-sensitive models are standardized inside a scikit-learn pipeline [29]. An RBF support-vector "
        "machine is omitted at this sample size for tractability. The hybrid model is evaluated on the same "
        "stratified 75/25 hold-out as the confusion analysis and is reported in Table 4.",
    ),
    (
        "Hardware. Experiments are CPU-only. The hybrid deep model, at 0.929 macro-F1, trails the boosted "
        "trees, a known pattern for tabular data, and was not GPU-tuned.",
        "Hardware. Experiments are CPU-only and were not GPU-tuned. The hybrid deep model trails boosted "
        "trees on this tabular task (Table 4), a known pattern for structured 802.11 features.",
    ),
    ("zero-day", "unseen-attack"),
    ("Zero-day", "Unseen-attack"),
    ("ZERO-DAY", "UNSEEN-ATTACK"),
    (
        "C4) Zero-day detection. Five unsupervised detectors are trained on benign traffic only, with an "
        "honest per-attack breakdown that includes attacks which are intrinsically hard to catch as anomalies.",
        "C4) Unseen-attack anomaly detection. Five unsupervised detectors are trained on benign traffic only "
        "and tested on held-out labelled attacks from the same AWID3 testbed (not out-of-distribution zero-day "
        "novelty), with an honest per-attack breakdown that includes attacks which are intrinsically hard to "
        "catch as anomalies.",
    ),
    (
        "That is what makes the result a fair proxy for a genuine zero-day.",
        "This measures unseen-attack detection within the AWID3 capture regime; it is not a claim of "
        "out-of-distribution zero-day generalisation.",
    ),
    (
        "Algorithm 3  Unsupervised zero-day detection with a calibrated threshold",
        "Algorithm 3  Unsupervised unseen-attack detection with a calibrated threshold",
    ),
    (
        "B. GENERALIZATION BEYOND AWID3",
        "B. PIPELINE PORTABILITY (SECONDARY)",
    ),
    (
        "The pipeline is not tied to AWID3. Its dataset-agnostic front end — automatic label resolution, "
        "identifier and constant-column removal, then the same benchmark suite — runs unchanged on other "
        "public intrusion corpora. We verified this on two flow-based datasets, CIC-IDS-2017 and "
        "CSE-CIC-IDS2018, both built by CICFlowMeter from wired network captures. These are not 802.11 "
        "traces, so the point is method transfer, not a like-for-like wireless comparison. On a stratified "
        "sample of each, under the same stratified cross-validation, boosted trees again lead (Table 7): "
        "0.982 macro-F1 on CIC-IDS-2017 and 0.902 on the larger and noisier CSE-CIC-IDS2018. The gap "
        "tracks the known difficulty of the two corpora and confirms that the leakage-aware, explainable "
        "workflow generalizes past the wireless dataset it was built for.",
        "The preprocessing front end — automatic label resolution, identifier and constant-column removal, "
        "then the same benchmark suite — is dataset-agnostic. As a secondary sanity check (not evidence for "
        "802.11 generalisation), we ran the same workflow on two wired flow corpora, CIC-IDS-2017 and "
        "CSE-CIC-IDS2018. These traces differ in modality, feature space and label semantics from AWID3; "
        "Table 7 therefore shows pipeline portability only. Boosted trees again lead (0.982 and 0.902 "
        "macro-F1), which supports reproducibility of the tooling but does not substitute for cross-wireless "
        "validation (e.g., AWID2) or live traffic.",
    ),
    (
        "which indicates the method transfers beyond 802.11.",
        "which indicates the tooling is portable, not that wireless conclusions transfer to wired flow data.",
    ),
    (
        "Section VII-B shows the pipeline transfers to wired flow datasets, but cross-wireless-dataset validation",
        "Section VII-B shows the pipeline runs on wired flow datasets as a portability check, but cross-wireless-dataset validation",
    ),
    (
        "TABLE 7. Cross-dataset generalization of the same pipeline (stratified CV). The CIC-IDS corpora are "
        "wired flow-based, included to show method transfer rather than as wireless results.",
        "TABLE 7. Pipeline portability on wired flow corpora (stratified CV). Not a wireless generalisation claim.",
    ),
    (
        "REPRODUCIBILITY",
        f"REPRODUCIBILITY\nSource code: {REPO_URL}\nFeature list: config/leakage_controlled_features.txt ({len(FEATURES)} fields).\n",
    ),
    (
        "Preprocessing, models, explainers and figure scripts are released, together with the configuration "
        "and the leakage-control feature policy, to enable independent verification.",
        "Preprocessing, models, explainers and figure scripts are released with the configuration, the "
        "leakage-control feature policy and the exact 34-feature list. A fixed random seed (42) and "
        "build_awid3.py reproduce the 74,270-row evaluation set.",
    ),
    (
        "III. DETECTION SYSTEM OVERVIEW",
        "III. EVALUATION CONTEXT",
    ),
    (
        "The evaluation sits inside a deployable, event-driven platform (Fig. 1). Distributed 802.11 sensors "
        "capture frames and stream them, tagged by sensor, to a central service that performs windowed feature "
        "extraction, model inference, explanation and policy-gated response. This paper focuses on the "
        "learning and evaluation parts. We describe the platform only to make clear that the models run in a "
        "realistic, real-time setting rather than in isolation.",
        "The offline experiments use the same feature schema and inference path as our companion streaming "
        "implementation: sensors buffer frames into tumbling windows, extract a numeric feature vector, and "
        "score it with the trained model. Fig. 1 sketches this path at a high level. Operational deployment "
        "details (message bus, storage backends, policy-gated response) are documented in the supplementary "
        "reproducibility package and are outside the scope of the learning evaluation reported here.",
    ),
    (
        "FIGURE 1. Detection platform: capture, streaming, feature extraction, inference, and explanation/response.",
        "FIGURE 1. High-level detection path: capture, windowed feature extraction, model inference, and optional explanation/response.",
    ),
    (
        "Leakage split, but not measured; no XAI/zero-day",
        "Leakage split, but not measured; no XAI/unseen-attack",
    ),
    (
        "INDEX TERMS Wireless intrusion detection, IEEE 802.11, AWID3, data leakage, explainable AI, SHAP, "
        "anomaly detection, zero-day, machine learning, network security.",
        "INDEX TERMS Wireless intrusion detection, IEEE 802.11, AWID3, data leakage, explainable AI, SHAP, "
        "anomaly detection, unseen-attack detection, machine learning, network security.",
    ),
]


def _set_paragraph_text(paragraph: Paragraph, new_text: str) -> None:
  if not paragraph.runs:
    paragraph.add_run(new_text)
    return
  paragraph.runs[0].text = new_text
  for run in paragraph.runs[1:]:
    run.text = ""


def replace_in_paragraph(paragraph: Paragraph) -> bool:
    text = paragraph.text
    if not text:
        return False
    new = text
    for old, repl in REPLACEMENTS:
        if old in new:
            new = new.replace(old, repl)
    if new != text:
        _set_paragraph_text(paragraph, new)
        return True
    return False


def replace_in_table(table: Table) -> int:
    n = 0
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if replace_in_paragraph(paragraph):
                    n += 1
    return n


def remove_algorithm2_paragraphs(doc: Document) -> None:
    """Drop Algorithm 2 block left after Section III shortening."""
    drop = False
    to_remove: list[Paragraph] = []
    for paragraph in doc.paragraphs:
        t = paragraph.text.strip()
        if t.startswith("Algorithm 2"):
            drop = True
        if drop:
            to_remove.append(paragraph)
            if t.startswith("IV. DATASET"):
                to_remove.pop()  # keep section header
                drop = False
    for paragraph in to_remove:
        p = paragraph._element
        p.getparent().remove(p)


def add_hybrid_table_row(doc: Document) -> None:
    for table in doc.tables:
        if not table.rows:
            continue
        header = " ".join(c.text for c in table.rows[0].cells)
        if "macro-F1" in header and "Lat." in header:
            for row in table.rows:
                if row.cells and row.cells[0].text.strip() == "Logistic Regression":
                    if any("Hybrid" in r.cells[0].text for r in table.rows):
                        return
                    new = table.add_row()
                    vals = [
                        "Hybrid (CNN+Transformer)",
                        "0.974",
                        "0.929",
                        "0.999",
                        "0.931",
                        "0.928",
                        "850",
                    ]
                    for cell, val in zip(new.cells, vals):
                        cell.text = val
                    note = doc.add_paragraph(
                        "Hybrid results are from the stratified 75/25 hold-out (not 5-fold CV)."
                    )
                    note.style = "Normal"
                    return


def add_feature_list_note(doc: Document) -> None:
    needle = "TABLE 2. Curated AWID3 evaluation set"
    for paragraph in doc.paragraphs:
        if needle in paragraph.text and "leakage_controlled_features" not in paragraph.text:
            paragraph.text = (
                paragraph.text.replace(
                    "34 leakage-controlled features.",
                    "34 leakage-controlled features (see config/leakage_controlled_features.txt).",
                )
            )
            return


def main() -> None:
    if not DOCX.exists():
        raise SystemExit(f"Missing {DOCX}")
    if not BACKUP.exists():
        shutil.copy2(DOCX, BACKUP)

    doc = Document(str(DOCX))
    changed = 0
    for paragraph in doc.paragraphs:
        if replace_in_paragraph(paragraph):
            changed += 1
    for table in doc.tables:
        changed += replace_in_table(table)

    remove_algorithm2_paragraphs(doc)
    add_hybrid_table_row(doc)
    add_feature_list_note(doc)

    doc.save(str(DOCX))
    print(f"Patched {changed} paragraph/cell blocks -> {DOCX}")


if __name__ == "__main__":
    main()
