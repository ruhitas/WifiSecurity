"""Remove the Ceylan et al. reference (prose, Table 1 row, reference entry)."""
from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "AWID3_Paper_IEEE_Access.docx"
BACKUP = ROOT / "AWID3_Paper_IEEE_Access_pre_remove_ceylan.docx"

OLD_PROSE = (
    "A very recent CNN pipeline is telling here [36]. It turns selected AWID3 features into 6x6 "
    "grayscale images and reports 99.93 percent on eight classes, and it does apply a capture-based "
    "split to hold leakage off. The split is where it stops. Leakage is prevented, not measured; there "
    "is no explanation, no unsupervised test for unseen attacks, and no latency budget. Our paper "
    "measures the leakage that such a split avoids, then adds the parts it leaves out."
)
NEW_PROSE = (
    "Even where a capture-based split is used to hold leakage off, the leakage is prevented rather than "
    "measured, and explanation, an unsupervised test for unseen attacks, and a latency budget are "
    "typically absent. Our paper measures the leakage that such a split avoids, then adds the parts it "
    "leaves out."
)


def set_text(p, text: str) -> None:
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def main() -> None:
    if not BACKUP.exists():
        shutil.copy2(DOCX, BACKUP)
    doc = Document(str(DOCX))

    # 1) prose rewrite (Section II)
    done = False
    for p in doc.paragraphs:
        if OLD_PROSE in p.text:
            set_text(p, p.text.replace(OLD_PROSE, NEW_PROSE, 1))
            done = True
            break
    if not done:
        raise RuntimeError("Ceylan prose block not found")

    # 2) remove Table 1 Ceylan row
    t = doc.tables[0]
    for row in list(t.rows):
        if any("Ceylan" in c.text for c in row.cells):
            row._element.getparent().remove(row._element)
            break
    else:
        raise RuntimeError("Ceylan row not found in Table 1")

    # 3) remove reference-list entry for Ceylan
    for p in list(doc.paragraphs):
        if re.match(r"^\[\d+\]", p.text.strip()) and "Ceylan" in p.text:
            p._element.getparent().remove(p._element)
            break
    else:
        raise RuntimeError("Ceylan reference entry not found")

    doc.save(str(DOCX))
    print("Removed Ceylan et al.: prose rewritten, Table 1 row deleted, reference entry deleted.")


if __name__ == "__main__":
    main()
