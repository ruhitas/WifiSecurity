# -*- coding: utf-8 -*-
"""Minor-revision round: novelty softening, cap rationale, weighted-F1, abstract details, antsignal note."""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "AWID3_Paper_IEEE_Access.docx"
BACKUP = ROOT / "AWID3_Paper_IEEE_Access_pre_minor_rev.docx"


def set_text(p, text: str) -> None:
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def edit(doc, anchor: str, old: str, new: str) -> None:
    for p in doc.paragraphs:
        if anchor in p.text and old in p.text:
            set_text(p, p.text.replace(old, new, 1))
            return
    raise RuntimeError(f"not found: {anchor!r} / {old!r}")


def main() -> None:
    if not BACKUP.exists():
        shutil.copy2(DOCX, BACKUP)
    doc = Document(str(DOCX))

    # 1) Soften "almost none measures leakage" and credit split-aware prior work
    edit(
        doc,
        "Almost none of them measures leakage",
        "Almost none of them measures leakage or attaches an explanation, which is the gap we take up.",
        "To be fair, a minority of studies already avoid the worst form of leakage by construction \u2014 "
        "capture-aware or session-aware splits appear in parts of this literature \u2014 but avoiding leakage "
        "and quantifying it are different things: none of these works reports how large the inflation would "
        "have been, and explanations attached to individual detections remain rare. Measuring the former and "
        "supplying the latter is the gap we take up.",
    )

    # 2) Justify the 20k/8k caps beyond imbalance
    edit(
        doc,
        "Sampling is capped and seeded:",
        "drawn with a fixed random seed of 42; classes that are smaller than their cap are taken in full.",
        "drawn with a fixed random seed of 42; classes that are smaller than their cap are taken in full. "
        "The caps themselves are pragmatic rather than tuned: 8,000 retains most attack classes in full "
        "(only the largest are subsampled), 20,000 keeps benign traffic the largest single class \u2014 about "
        "27 percent of rows, a realistic majority \u2014 and the resulting 74,270 rows keep the full nine-model, "
        "five-fold protocol tractable on a CPU workstation.",
    )

    # 3) Support-weighted F1 next to the main-table statistical summary
    edit(
        doc,
        "Statistical summary (XGBoost, best model):",
        "ROC-AUC CI = [0.9996, 0.9998].",
        "ROC-AUC CI = [0.9996, 0.9998]; support-weighted F1 = 0.985 (hold-out).",
    )

    # 4) Abstract: dataset size, bootstrap CI, explicit window size
    edit(
        doc,
        "ABSTRACT Wi-Fi intrusion detectors",
        "On the author-curated subset, and after stripping every temporal, sequence and address identifier, "
        "a gradient-boosted model reaches 0.976 macro-F1 and 0.9998 ROC-AUC across fourteen classes.",
        "On the author-curated 74,270-sample subset, and after stripping every temporal, sequence and address "
        "identifier, a gradient-boosted model reaches 0.976 macro-F1 (bootstrap 95% CI 0.966\u20130.979) and "
        "0.9998 ROC-AUC across fourteen classes.",
    )
    edit(
        doc,
        "ABSTRACT Wi-Fi intrusion detectors",
        "a full detection stays well under a millisecond per window",
        "a full detection stays well under a millisecond per 40-frame window",
    )

    # 5) Introduction: why 14-class granularity is more realistic
    edit(
        doc,
        "This paper takes the opposite stance.",
        "low-latency detection lives comfortably inside that protocol.",
        "low-latency detection lives comfortably inside that protocol. We also evaluate at full fourteen-class "
        "granularity rather than the coarse three-to-five-class setups common in the literature: an operator "
        "must know which attack is under way to respond, and coarse grouping hides exactly the confusions "
        "\u2014 deauthentication versus disassociation, for example \u2014 that decide operational usefulness.",
    )

    # 6) Appendix A: radiotap.dbm_antsignal after leakage control (backed by permutation data)
    edit(
        doc,
        "The count is not hand-tuned:",
        "yields exactly 34.",
        "yields exactly 34. One retained field deserves comment: radiotap.dbm_antsignal is the strongest "
        "leakage carrier under the naive split (permutation importance 0.23, the single largest), yet it "
        "remains the top predictor on the curated same-capture data (0.34). The two roles differ: across "
        "captures it fingerprints the recording session, whereas within a capture it reflects legitimate "
        "physical variation between stations, which is exactly the signal a live detector would also see.",
    )

    doc.save(str(DOCX))
    print("Applied 6 minor-revision edits.")


if __name__ == "__main__":
    main()
