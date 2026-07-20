"""Renumber algorithms to order-of-appearance and add the missing Algorithm reference."""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "AWID3_Paper_IEEE_Access.docx"
BACKUP = ROOT / "AWID3_Paper_IEEE_Access_pre_algorithms.docx"


def set_text(p, text: str) -> None:
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def main() -> None:
    if not BACKUP.exists():
        shutil.copy2(DOCX, BACKUP)
    doc = Document(str(DOCX))

    # 1) swap caption numbers: detection (was Alg2) -> Alg1 ; evaluation (was Alg1) -> Alg2
    for t in doc.tables:
        cell0 = t.rows[0].cells[0]
        head = cell0.paragraphs[0]
        if head.text.startswith("Algorithm 2  Real-time"):
            set_text(head, head.text.replace("Algorithm 2  Real-time", "Algorithm 1  Real-time", 1))
        elif head.text.startswith("Algorithm 1  Leakage-controlled"):
            set_text(head, head.text.replace("Algorithm 1  Leakage-controlled", "Algorithm 2  Leakage-controlled", 1))

    # 2) prose that referenced the evaluation algorithm (now Algorithm 2)
    for p in doc.paragraphs:
        if "Algorithm 1 is the recipe our code actually runs" in p.text:
            set_text(p, p.text.replace(
                "Algorithm 1 is the recipe our code actually runs",
                "Algorithm 2 is the recipe our code actually runs", 1))
            break

    # 3) add the missing reference to the real-time detection algorithm (now Algorithm 1)
    for p in doc.paragraphs:
        if "Figure 1 sketches this path at a high level." in p.text:
            set_text(p, p.text.replace(
                "Figure 1 sketches this path at a high level.",
                "Figure 1 sketches this path at a high level, and Algorithm 1 states the streaming "
                "detection loop (windowing, threshold-gated alerting, and policy-gated response) in full.", 1))
            break

    doc.save(str(DOCX))
    print("Algorithms renumbered to appearance order; Algorithm 1 now referenced in Section III.")


if __name__ == "__main__":
    main()
