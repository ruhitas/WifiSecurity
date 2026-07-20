"""Reviewer round 2: sharpen novelty, reproducibility, minority-class, threshold, edge deploy."""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "AWID3_Paper_IEEE_Access.docx"
BACKUP = ROOT / "AWID3_Paper_IEEE_Access_pre_reviewer2.docx"


def set_text(p, text: str) -> None:
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def edit(doc, anchor: str, old: str, new: str) -> None:
    for p in doc.paragraphs:
        if anchor in p.text and old in p.text:
            set_text(p, p.text.replace(old, new, 1))
            return
    raise RuntimeError(f"anchor/old not found: {anchor!r}")


def main() -> None:
    if not BACKUP.exists():
        shutil.copy2(DOCX, BACKUP)
    doc = Document(str(DOCX))

    # 1) C1 contribution: add leakage source attribution
    edit(
        doc,
        "C1) Leakage measurement.",
        "which isolates the cause as capture-environment leakage rather than a few identifying columns.",
        "which isolates the cause as capture-environment leakage rather than a few identifying columns. "
        "Going beyond detection, we attribute the source: a leave-one-feature-group-out analysis (Figure 11) "
        "shows that radio-layer fields, not attack behaviour, carry most of the inflated signal.",
    )

    # 2) Table 1 fairness vs recent SOTA
    edit(
        doc,
        "Table 1 places representative results next to ours.",
        "The setups differ, so it is indicative, not a ranking, and we say so plainly.",
        "The setups differ, so it is indicative, not a ranking, and we say so plainly. The most recent "
        "temporal-context [20] and hybrid CNN-LSTM [24] models report strong headline numbers, but on different "
        "class sets and without a measured leakage split; comparing our leakage-controlled score against "
        "potentially leakage-inflated ones would reproduce the very confound this paper isolates, so we decline "
        "that comparison rather than overstate it.",
    )

    # 3) Feature-selection code reference
    edit(
        doc,
        "For feature selection we compare mutual information",
        "then form a consensus ranking and report a feature-count ablation.",
        "then form a consensus ranking and report a feature-count ablation (the six selectors and the "
        "average-rank consensus rule are released in src/wids/feature_selection/selectors.py).",
    )

    # 4) Reproducibility summary
    edit(
        doc,
        "quantified against the naive across-capture split.",
        "quantified against the naive across-capture split.",
        "quantified against the naive across-capture split. For reproducibility, every estimator uses a fixed "
        "seed of 42, tree ensembles use 300 trees without early stopping, and both the five-fold split and the "
        "75/25 hold-out are seeded; the complete configuration is released in src/wids/benchmark/runner.py.",
    )

    # 5) Threshold / FPR robustness caveat
    edit(
        doc,
        "Algorithm 3 never lets an attack sample touch training.",
        "That is what makes the result a fair proxy for a genuine unseen-attack.",
        "That is what makes the result a fair proxy for a genuine unseen-attack. The operating threshold is a "
        "quantile of benign scores at the five-percent budget; ROC-AUC, our headline unseen-attack metric, is "
        "threshold-independent, while the stability of that operating point across different benign captures is "
        "left to future work.",
    )

    # 6) Minority-class reliability mitigation
    edit(
        doc,
        "Class imbalance.",
        "which limits confidence in their per-class metrics.",
        "which limits confidence in their per-class metrics. We therefore lead with macro-averaged F1, report "
        "bootstrap confidence intervals (Section VI), and use stratified folds that keep every class present in "
        "each partition; even so, per-class scores for the smallest categories (SQL_Injection, RogueAP) should be "
        "read as indicative, and larger same-capture samples are needed to tighten them.",
    )

    # 7) Full-archive future work strengthening
    edit(
        doc,
        "Full-archive protocol.",
        "is identified but not implemented here.",
        "is identified but not implemented here; it is the natural way to use the corpus scale without "
        "reintroducing leakage and is a priority extension alongside cross-wireless-dataset validation.",
    )

    # 8) Conclusion: formalize wording + add edge-deployment future work
    edit(
        doc,
        "We studied machine-learning intrusion detection",
        "Under a controlled protocol the honest picture emerges:",
        "Under a controlled protocol a rigorous, leakage-controlled picture emerges:",
    )
    edit(
        doc,
        "Next we will sample the full archive within captures",
        "and test the whole system on real deployments.",
        "test the whole system on real deployments; and compress the model for edge devices through ONNX export "
        "and post-training quantization.",
    )

    doc.save(str(DOCX))
    print("Applied 9 reviewer-round-2 edits.")


if __name__ == "__main__":
    main()
