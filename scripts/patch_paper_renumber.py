"""Renumber all citations to IEEE order-of-first-appearance and reorder REFERENCES."""
from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "AWID3_Paper_IEEE_Access.docx"
BACKUP = ROOT / "AWID3_Paper_IEEE_Access_pre_renumber.docx"


def para_text(el) -> str:
    return "".join(t.text or "" for t in el.iter(W + "t"))


def set_text(p, text: str) -> None:
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def main() -> None:
    if not BACKUP.exists():
        shutil.copy2(DOCX, BACKUP)

    doc = Document(str(DOCX))
    body = doc.element.body

    # 1) first-appearance order (body reading order, before REFERENCES)
    order: list[int] = []
    seen: set[int] = set()
    stop = False
    for child in body:
        tag = child.tag.split("}")[-1]
        if tag not in ("p", "tbl"):
            continue
        txt = para_text(child)
        if tag == "p" and txt.strip() == "REFERENCES":
            stop = True
            continue
        if stop:
            continue
        for m in re.finditer(r"\[(\d+)\]", txt):
            n = int(m.group(1))
            if n not in seen:
                seen.add(n)
                order.append(n)

    old_to_new = {old: i + 1 for i, old in enumerate(order)}

    # 2) capture + remove existing reference-definition paragraphs
    ref_texts: dict[int, str] = {}
    to_remove = []
    for p in doc.paragraphs:
        m = re.match(r"^\[(\d+)\]", p.text.strip())
        if m:
            ref_texts[int(m.group(1))] = p.text.strip()
            to_remove.append(p)
    for p in to_remove:
        p._element.getparent().remove(p._element)

    # 3) renumber in-text citations (defs already removed) via placeholder pass
    def remap(text: str) -> str:
        text = re.sub(r"\[(\d+)\]", lambda m: f"\x00{old_to_new[int(m.group(1))]}\x01", text)
        return text.replace("\x00", "[").replace("\x01", "]")

    for p in doc.paragraphs:
        if "[" in p.text:
            set_text(p, remap(p.text))
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if "[" in p.text:
                        set_text(p, remap(p.text))

    # 4) rebuild ordered reference list after REFERENCES heading
    ref_heading = next(p for p in doc.paragraphs if p.text.strip() == "REFERENCES")
    ref_el = ref_heading._element
    new_lines = []
    for new_num in range(len(order), 0, -1):
        old_num = order[new_num - 1]
        text = re.sub(r"^\[\d+\]", f"[{new_num}]", ref_texts[old_num])
        new_lines.append(text)
    for line in new_lines:  # already reversed range -> insert via addnext keeps order
        np = doc.add_paragraph(line)
        ref_el.addnext(np._element)

    doc.save(str(DOCX))
    print("Renumbered", len(order), "references to first-appearance order.")
    print("map sample:", {k: old_to_new[k] for k in sorted(old_to_new)})


if __name__ == "__main__":
    main()
