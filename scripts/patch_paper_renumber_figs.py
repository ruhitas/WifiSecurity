"""Renumber figure references/captions to IEEE order-of-first-appearance."""
from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "AWID3_Paper_IEEE_Access.docx"
BACKUP = ROOT / "AWID3_Paper_IEEE_Access_pre_renumber_figs.docx"

FIG_RE = re.compile(
    r"((?:FIGURE|Figures|Figure)\s+)(\d+)((?:\s*(?:-|\u2013|to|and|,)\s*)(\d+))?"
)


def para_text(el) -> str:
    return "".join(t.text or "" for t in el.iter(W + "t"))


def set_text(p, text: str) -> None:
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def build_order(doc: Document) -> list[int]:
    order: list[int] = []
    seen: set[int] = set()
    for child in doc.element.body:
        if child.tag.split("}")[-1] not in ("p", "tbl"):
            continue
        txt = para_text(child)
        for m in FIG_RE.finditer(txt):
            a = int(m.group(2))
            b = int(m.group(4)) if m.group(4) else None
            nums = list(range(a, b + 1)) if b else [a]
            for n in nums:
                if n not in seen:
                    seen.add(n)
                    order.append(n)
    return order


def main() -> None:
    if not BACKUP.exists():
        shutil.copy2(DOCX, BACKUP)
    doc = Document(str(DOCX))

    order = build_order(doc)
    mapping = {old: i + 1 for i, old in enumerate(order)}
    if order == list(range(1, len(order) + 1)):
        print("Figures already sequential; nothing to do.")
        return

    def remap(text: str) -> str:
        def repl(m):
            head = m.group(1)
            a = int(m.group(2))
            na = mapping[a]
            if m.group(4):
                conn = m.group(3)[: m.group(3).rfind(m.group(4))]
                b = int(m.group(4))
                return f"{head}\x00{na}\x01{conn}\x00{mapping[b]}\x01"
            return f"{head}\x00{na}\x01"

        text = FIG_RE.sub(repl, text)
        return text.replace("\x00", "").replace("\x01", "")

    for p in doc.paragraphs:
        if FIG_RE.search(p.text):
            set_text(p, remap(p.text))
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if FIG_RE.search(p.text):
                        set_text(p, remap(p.text))

    doc.save(str(DOCX))
    print("Figure order was:", order)
    print("map old->new:", {k: mapping[k] for k in sorted(mapping)})


if __name__ == "__main__":
    main()
