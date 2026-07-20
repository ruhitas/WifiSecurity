"""Second-pass cleanup for AWID3_Paper_IEEE_Access.docx after review fixes."""
from __future__ import annotations

from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "AWID3_Paper_IEEE_Access.docx"
FEATURES_FILE = "config/leakage_controlled_features.txt"


def _set_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def main() -> None:
    doc = Document(str(DOCX))

    for paragraph in doc.paragraphs:
        t = paragraph.text
        if "Algorithm 2" in t:
            _set_text(paragraph, "")
        if "forty clean features" in t:
            _set_text(paragraph, t.replace("forty clean features", "thirty-four clean features"))
        if t.startswith("REPRODUCIBILITY"):
            _set_text(
                paragraph,
                "REPRODUCIBILITY\n"
                "Source code: https://github.com/ruhitas/awid3-leakage-aware-ids\n"
                f"Feature list: {FEATURES_FILE} (34 fields).\n"
                "Preprocessing, models, explainers and figure scripts are released with the "
                "configuration, the leakage-control feature policy and the exact 34-feature list. "
                "A fixed random seed (42) and build_awid3.py reproduce the 74,270-row evaluation set.",
            )

    for paragraph in doc.paragraphs:
        if paragraph.text.startswith("Preprocessing, models, explainers and figure scripts are released with"):
            if paragraph.text.startswith("REPRODUCIBILITY"):
                continue
            _set_text(paragraph, "")

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if cells == ["Full archive, restricted to clean features", "40", "~1.00", "Leakage persists — not feature-driven"]:
                row.cells[1].text = "34"
            if cells == ["Curated subset (same-capture interleave)", "40", "0.976", "Honest, leakage-controlled baseline"]:
                row.cells[1].text = "34"

    doc.save(str(DOCX))
    print(f"Cleanup complete -> {DOCX}")


if __name__ == "__main__":
    main()
