"""Add missing Wi-Fi IDS references, fix citation order, and restore REFERENCES layout."""
from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "AWID3_Paper_IEEE_Access.docx"
BACKUP = ROOT / "AWID3_Paper_IEEE_Access_pre_refs.docx"

NEW_REFS = [
    "[34] D. S. Berman, A. L. Buczak, J. S. Chavis, and C. L. Corbett, "
    '"A survey of deep learning methods for cyber security," Information, vol. 10, no. 4, art. 122, 2019.',
    "[35] A. M. Mahfouz, D. Venugopal, and S. G. Shiva, "
    '"Comparative analysis of ML classifiers for network intrusion detection," '
    "in Proc. 4th Int. Congr. Inf. Commun. Technol., Springer, 2020, pp. 193\u2013207.",
    "[36] R. Kohavi and G. H. John, "
    '"Wrappers for feature subset selection," Artif. Intell., vol. 97, nos. 1\u20132, pp. 273\u2013324, 1997.',
    "[37] S. Bhandari, A. K. Kukreja, A. Lazar, A. Sim, and K. Wu, "
    '"Feature selection improves tree-based classification for wireless intrusion detection," '
    "in Proc. 3rd Int. Workshop Syst. Netw. Telemetry Analytics, 2020, pp. 19\u201326.",
    "[38] Y. Qin, B. Li, M. Yang, and Z. Yan, "
    '"Attack detection for wireless enterprise network: A machine learning approach," '
    "in Proc. IEEE Int. Conf. Signal Process., Commun. Comput. (ICSPCC), 2018, pp. 1\u20136.",
    "[39] J.-h. Woo, J.-Y. Song, and Y.-J. Choi, "
    '"Performance enhancement of deep neural network using feature selection and preprocessing for intrusion detection," '
    "in Proc. Int. Conf. Artif. Intell. Inf. Commun. (ICAIIC), IEEE, 2019, pp. 415\u2013417.",
    "[40] M. Agarwal, S. Biswas, and S. Nandi, "
    '"Detection of de-authentication DoS attacks in Wi-Fi networks: A machine learning approach," '
    "in Proc. IEEE Int. Conf. Syst., Man, Cybern., 2015, pp. 246\u2013251.",
    "[41] M. Agarwal, S. Purwar, S. Biswas, and S. Nandi, "
    '"Intrusion detection system for PS-Poll DoS attack in 802.11 networks using real time discrete event system," '
    "IEEE/CAA J. Automatica Sinica, vol. 4, no. 4, pp. 792\u2013808, 2017.",
    "[42] R. Abdulhammed, M. Faezipour, A. Abuzneid, and A. Alessa, "
    '"Effective features selection and machine learning classifiers for improved wireless intrusion detection," '
    "in Proc. Int. Symp. Netw., Comput. Commun. (ISNCC), IEEE, 2018, pp. 1\u20136.",
    "[43] M. E. Aminanto, H. C. Tanuwidjaja, P. D. Yoo, and K. Kim, "
    '"Wi-Fi intrusion detection using weighted-feature selection for neural networks classifier," '
    "in Proc. Int. Workshop Big Data Inf. Security (IWBIS), IEEE, 2017, pp. 99\u2013104.",
]


def set_text(p, text: str) -> None:
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def replace_all(doc: Document, old: str, new: str) -> None:
    for p in doc.paragraphs:
        if old in p.text:
            set_text(p, p.text.replace(old, new))
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if old in p.text:
                        set_text(p, p.text.replace(old, new))


def rebuild_reference_section(doc: Document) -> None:
    existing: list[str] = []
    to_remove = []
    for p in doc.paragraphs:
        m = re.match(r"^\[(\d+)\]", p.text.strip())
        if m:
            existing.append(p.text.strip())
            to_remove.append(p)

    for p in to_remove:
        p._element.getparent().remove(p._element)

    ref_heading = next(p for p in doc.paragraphs if p.text.strip() == "REFERENCES")
    ref_el = ref_heading._element

    all_refs = existing + NEW_REFS
    for line in reversed(all_refs):
        np = doc.add_paragraph(line)
        ref_el.addnext(np._element)

    seen_caption = False
    for p in list(doc.paragraphs):
        if p.text.strip().startswith("TABLE A1. Thirty-four leakage-controlled"):
            if seen_caption:
                p._element.getparent().remove(p._element)
            else:
                seen_caption = True


def main() -> None:
    if not BACKUP.exists():
        shutil.copy2(DOCX, BACKUP)

    doc = Document(str(DOCX))

    replace_all(
        doc,
        "The list of practical attacks is long: deauthentication and disassociation floods, key reinstallation (KRACK) [1]",
        "The list of practical attacks is long: deauthentication and disassociation floods [40], [41], "
        "key reinstallation (KRACK) [1]",
    )

    replace_all(
        doc,
        "AWID3 extended it to enterprise WPA2/WPA3 settings and added attacks such as KRACK and Kr00k [3]. "
        "Work built on these datasets is plentiful. Expert feature selection",
        "AWID3 extended it to enterprise WPA2/WPA3 settings and added attacks such as KRACK and Kr00k [3]. "
        "Work built on these datasets is plentiful; surveys of deep learning for cyber security [34] and "
        "comparative ML benchmarks on public intrusion traces [35] frame the wider literature. "
        "Expert feature selection",
    )

    replace_all(
        doc,
        "filter-based feature engineering has fed deep classifiers [16], and semi-supervised models have been tried "
        "for 802.11 anomaly detection [17].",
        "filter-based feature engineering has fed deep classifiers [16]. Tree-based and wrapper-style feature "
        "selection has sharpened wireless IDS accuracy [36], [37], enterprise-traffic classifiers have been reported "
        "[38], and preprocessing-aware deep pipelines [39] together with multi-algorithm screening [42] extend "
        "the same line; weighted neural selection [43] foreshadowed later impersonation-focused designs. "
        "Semi-supervised models have been tried for 802.11 anomaly detection [17].",
    )

    replace_all(
        doc,
        "gradient boosting, a multilayer perceptron, XGBoost [27] and LightGBM [28] \u2014 plus a tenth hybrid deep model "
        "that fuses a one-dimensional convolutional branch with a Transformer self-attention branch [30]. "
        "Scale-sensitive models are standardized inside a scikit-learn pipeline [29].",
        "gradient boosting, a multilayer perceptron, XGBoost [27] and LightGBM [28]. Scale-sensitive models are "
        "standardized inside a scikit-learn pipeline [29] \u2014 plus a tenth hybrid deep model that fuses a "
        "one-dimensional convolutional branch with a Transformer self-attention branch [30].",
    )

    replace_all(
        doc,
        "L1-regularized logistic regression and recursive elimination, then form a consensus ranking",
        "L1-regularized logistic regression and recursive elimination (wrapper-style subset search [36]), "
        "then form a consensus ranking",
    )

    rebuild_reference_section(doc)
    doc.save(str(DOCX))
    print(f"Updated {DOCX.name}: {len(NEW_REFS)} references added (now [1]\u2013[43]); REFERENCES section restored.")


if __name__ == "__main__":
    main()
