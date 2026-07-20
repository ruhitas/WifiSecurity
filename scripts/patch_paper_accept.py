# -*- coding: utf-8 -*-
"""Acceptance-raising edits: WPA4, weighted-F1, author-curated clarity, signature caveat, platform genericization."""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "AWID3_Paper_IEEE_Access.docx"
BACKUP = ROOT / "AWID3_Paper_IEEE_Access_pre_accept.docx"


def set_text(p, text: str) -> None:
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def replace_para(doc, anchor: str, old: str, new: str) -> None:
    for p in doc.paragraphs:
        if anchor in p.text and old in p.text:
            set_text(p, p.text.replace(old, new, 1))
            return
    raise RuntimeError(f"not found: {anchor!r} / {old!r}")


def replace_cell(doc, old: str, new: str) -> None:
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if old in p.text:
                        set_text(p, p.text.replace(old, new, 1))
                        return
    raise RuntimeError(f"cell not found: {old!r}")


def main() -> None:
    if not BACKUP.exists():
        shutil.copy2(DOCX, BACKUP)
    doc = Document(str(DOCX))

    # m7 — WPA4 is not a standard; hedge it
    replace_para(
        doc,
        "The third is the moving target of the standard.",
        "WPA3 is still rolling out and WPA4 will follow; each brings new handshakes and new attacks",
        "WPA3 is still rolling out and future amendments to the standard will follow; each brings new "
        "handshakes and potentially new attacks",
    )

    # m5 — report support-weighted F1 alongside macro-F1
    replace_para(
        doc,
        "We bootstrap both macro-F1 and macro ROC-AUC",
        "Macro-F1 is 0.9728 with 95% CI [0.9658, 0.9793];",
        "Macro-F1 is 0.9728 with 95% CI [0.9658, 0.9793], and the support-weighted F1 is 0.985 \u2014 the gap "
        "confirms that the smallest attack classes, not broad error, hold the macro figure down;",
    )

    # m3 — disambiguate "author-curated" (dataset creators, not us)
    replace_para(
        doc,
        "The subset is not hand-picked",
        "We start from the author-curated AWID3 reduced release,",
        "We start from the reduced per-attack CSV release published by the AWID3 authors [3] (curated by the "
        "dataset creators, not by us),",
    )

    # m4 — flag signature-like behaviour for port-driven classes
    replace_para(
        doc,
        "SHAP reasoning reports make each verdict auditable.",
        "Explanation adds operator value without contradicting the numbers.",
        "Because some of these cues \u2014 destination ports in particular \u2014 are close to what a signature "
        "rule would inspect, for those classes the model is effectively doing signature-like matching rather "
        "than learning novel behaviour; we note this instead of overclaiming learned semantics. Explanation "
        "adds operator value without contradicting the numbers.",
    )

    # M7 — remove product-specific names from the deployment algorithm (reads as advertising)
    replace_cell(doc, "cache v in Redis; publish v to Kafka",
                 "cache v; publish v to the streaming bus")
    replace_cell(doc, "persist event to MSSQL; index alert in Elasticsearch",
                 "persist the event; index the alert for search")

    doc.save(str(DOCX))
    print("Applied acceptance-raising edits: WPA4, weighted-F1, author-curated, signature caveat, platform genericization.")


if __name__ == "__main__":
    main()
