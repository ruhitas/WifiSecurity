"""Create TurkceOzet.docx — plain-language Turkish summary of the AWID3 paper."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

OUT = Path(__file__).resolve().parents[1] / "TurkceOzet.docx"


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    return p


def h2(doc, text):
    return doc.add_heading(text, level=2)


def p(doc, text, bold=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    return para


def bullet(doc, text):
    para = doc.add_paragraph(text, style="List Bullet")
    for run in para.runs:
        run.font.size = Pt(11)
        run.font.name = "Calibri"
    return para


def main():
    doc = Document()

    title = doc.add_heading(
        "Wi-Fi Güvenlik Tespit Sistemi — Türkçe Özet", level=0
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(
        "Makale: Leakage-Aware and Explainable Machine Learning for IEEE 802.11 "
        "Intrusion Detection (AWID3)\n"
        "Yazar: Ruhi Taş — OSTİM Teknik Üniversitesi"
    )
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # ---- 1 ----
    h1(doc, "1. Bu çalışma ne hakkında?")
    p(
        doc,
        "Bu çalışma, Wi-Fi (kablosuz ağ) saldırılarını otomatik tespit eden yapay zekâ "
        "sistemleriyle ilgilidir. Akademide ve sektörde sıkça kullanılan AWID3 adlı "
        "büyük bir Wi-Fi saldırı veri seti üzerinde, mevcut sonuçların ne kadar "
        "güvenilir olduğunu inceliyor.",
    )
    p(
        doc,
        "Kısaca sorulan soru şu: “Wi-Fi saldırı tespiti yapan modeller gerçekten saldırıyı "
        "mı öğreniyor, yoksa veri setindeki bir hile sayesinde mi yüksek skor alıyor?”",
    )

    # ---- 2 ----
    h1(doc, "2. Temel sorun: “Sızıntı” (leakage)")
    p(
        doc,
        "Birçok bilimsel makale AWID3 üzerinde %99’un üzerinde başarı bildiriyor. "
        "Bu çalışma gösteriyor ki bu rakamlar bazen yanıltıcı olabiliyor.",
    )
    p(doc, "Neden?", bold=True)
    p(
        doc,
        "Çünkü saldırı içeren kayıtlar ile normal trafik kayıtları çoğu zaman farklı "
        "kayıt oturumlarından geliyor. Model, saldırının kendisini değil; o oturumun "
        "özelliklerini (sinyal gücü, kanal, zaman damgası gibi) öğreniyor. Yani "
        "“saldırı var” diye değil, “bu kayıt başka bir oturumdan” diye ayırt ediyor.",
    )
    p(
        doc,
        "Bu etkiye capture-environment leakage (kayıt ortamı sızıntısı) deniyor. "
        "Sonuç: Model kâğıt üzerinde mükemmel görünüyor; gerçek hayatta ise zayıf kalabilir.",
    )

    # ---- 3 ----
    h1(doc, "3. Ne yaptık? (Basit adımlar)")
    h2(doc, "Adım A — Sızıntıyı ölçmek")
    bullet(
        doc,
        "AWID3’ün tüm arşivinden “naif” (dikkatsiz) örnekleme yaptık: saldırı ve normal "
        "trafik farklı oturumlardan geliyordu.",
    )
    bullet(
        doc,
        "Sonuç: Hemen her model, hatta basit doğrusal model bile yaklaşık 1,00 (mükemmel) "
        "skor aldı. Bu, sızıntının açık kanıtı.",
    )
    bullet(
        doc,
        "Hangi alanların sızıntıya yol açtığını da ölçtük: özellikle radyo sinyali "
        "(RSSI), kanal ve benzeri alanlar skorları şişiriyor.",
    )

    h2(doc, "Adım B — Dürüst bir test ortamı kurmak")
    bullet(
        doc,
        "Aynı kayıt oturumundan gelen normal ve saldırı çerçevelerini birlikte kullandık.",
    )
    bullet(
        doc,
        "Zaman damgası, MAC/IP adresi, sıra numarası gibi kimlik/oturum bilgilerini çıkardık.",
    )
    bullet(
        doc,
        "Geriye saldırı davranışını anlatan 34 özellik kaldı (liste makale ekinde).",
    )
    bullet(
        doc,
        "74.270 örnek, 14 sınıf (Normal + 13 saldırı türü) ile yeniden test ettik.",
    )

    h2(doc, "Adım C — Dürüst sonuçları raporlamak")
    bullet(
        doc,
        "En iyi model (XGBoost): yaklaşık %97,6 macro-F1 ve çok yüksek ROC-AUC. "
        "Bu, sızıntı kontrolünden sonra hâlâ güçlü bir sonuç.",
    )
    bullet(
        doc,
        "Hatalar mantıklı: Deauthentication ile Disassociation birbirine karışıyor "
        "(ikisi de benzer yönetim çerçevesi saldırıları).",
    )
    bullet(
        doc,
        "Sadece iyi bilinen saldırıları değil; modelin hiç görmediği saldırıları da "
        "(yalnızca normal trafik üzerinde eğitilmiş otoenkoder ile) yakalamayı denedik: "
        "yaklaşık 0,991 ROC-AUC, %5 civarı yanlış alarm.",
    )

    h2(doc, "Adım D — Açıklanabilirlik ve hız")
    bullet(
        doc,
        "SHAP ile her uyarı için “model neden böyle karar verdi?” sorusuna özellik "
        "bazlı açıklama üretildi (ör. SSDP için UDP portu, KRACK için kanal/çerçeve alanları).",
    )
    bullet(
        doc,
        "15–20 özellik yeterli; tam tespit penceresi milisaniyenin altında. "
        "Asıl süre modelden değil, özellik çıkarmadan geliyor.",
    )

    # ---- 4 ----
    h1(doc, "4. Son kullanıcıya ne ifade eder?")
    p(
        doc,
        "Bu makale bir “ürün broşürü” değil; bilimsel bir güvenilirlik uyarısı ve "
        "doğru ölçüm rehberidir. Pratikte şu anlama gelir:",
    )
    bullet(
        doc,
        "Wi-Fi saldırı tespit ürünü veya araştırması görüyorsanız, yalnızca “%99 doğruluk” "
        "rakamına bakmayın; veri nasıl bölünmüş, kimlik/zaman alanları çıkarılmış mı, "
        "sızıntı kontrol edilmiş mi sorun.",
    )
    bullet(
        doc,
        "Dürüst testte skor biraz düşebilir (ör. 1,00 → 0,976); bu kötü değil — "
        "gerçeğe daha yakın demektir.",
    )
    bullet(
        doc,
        "Açıklanabilir uyarılar (SHAP), güvenlik operatörünün “neden alarm çaldı?” "
        "sorusuna yanıt vermesine yardımcı olur.",
    )
    bullet(
        doc,
        "Sistem, canlı ağ için tasarlanmış bir pipeline ile uyumlu düşünülmüştür; "
        "makaledeki ana sonuçlar ise çevrimdışı, tekrarlanabilir deneylerdir.",
    )

    # ---- 5 ----
    h1(doc, "5. Ana sonuçlar (kısa tablo)")
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Konu"
    hdr[1].text = "Sonuç (anlaşılır dilde)"

    rows = [
        (
            "Naif (sızıntılı) ölçüm",
            "Neredeyse mükemmel skor (~1,00) — yanıltıcı olabilir",
        ),
        (
            "Sızıntı kontrollü ölçüm",
            "Dürüst skor ≈ 0,976 macro-F1 (14 sınıf)",
        ),
        (
            "Sızıntıyı taşıyan alanlar",
            "Özellikle radyo/kanal/sinyal alanları; kimlik/zaman alanları da riskli",
        ),
        (
            "Kullanılan özellik sayısı",
            "34 davranışsal özellik (254’ten politika + sabit sütun temizliği)",
        ),
        (
            "Görülmemiş saldırı (anomali)",
            "Otoenkoder ≈ 0,991 ROC-AUC, ~%5 yanlış pozitif",
        ),
        (
            "Açıklama",
            "SHAP ile özellik bazlı gerekçe raporları",
        ),
        (
            "Hız",
            "Pencere başına < 1 ms; model çok hızlı, özellik çıkarma baskın",
        ),
        (
            "Tekrarlanabilirlik",
            "Kod, özellik listesi ve pipeline açık olarak paylaşılıyor",
        ),
    ]
    for a, b in rows:
        cells = table.add_row().cells
        cells[0].text = a
        cells[1].text = b

    # ---- 6 ----
    h1(doc, "6. Ne yapılmadı / sınırlar (dürüstçe)")
    bullet(
        doc,
        "AWID2 gibi ikinci bir kablosuz veri setinde aynı sızıntı testi henüz yok "
        "(öncelikli gelecek iş).",
    )
    bullet(
        doc,
        "Canlı kurumsal ağda uzun süreli saha doğrulaması makalenin ana konusu değil; "
        "odak metodoloji ve dürüst kıyas.",
    )
    bullet(
        doc,
        "Derin hibrit model (CNN+Transformer) ek olarak raporlandı; ana tabloda "
        "ağaç tabanlı modellerle aynı 5-katlı çapraz doğrulama protokolünde değil "
        "(ek bölümde ayrı tutuldu).",
    )
    bullet(
        doc,
        "CIC-IDS gibi kablolu veri setleri yalnızca “aynı araçlar başka veriye de "
        "uygulanır” kontrolü içindir; Wi-Fi sonucu sayılmaz.",
    )

    # ---- 7 ----
    h1(doc, "7. Tek cümlelik özet")
    p(
        doc,
        "Wi-Fi saldırı tespitinde yüksek skorlar bazen “saldırıyı öğrenmekten” değil, "
        "“kayıt oturumunu ezberlemekten” gelir; sızıntıyı ölçüp temizledikten sonra "
        "hâlâ güçlü, açıklanabilir ve hızlı bir tespit mümkündür — ama dürüst "
        "protokol olmadan yayınlanan %99’lar yanıltıcı olabilir.",
        bold=True,
    )

    # ---- 8 ----
    h1(doc, "8. İlgili dosyalar")
    bullet(doc, "Makale: AWID3_Paper_IEEE_Access.docx")
    bullet(doc, "Özellik listesi: config/leakage_controlled_features.txt")
    bullet(doc, "Kanıt raporu: data/reports/reviewer_evidence.json")
    bullet(doc, "Kod paketi: github/ klasörü (reproducibility)")

    footer = doc.add_paragraph()
    footer.add_run(
        "\nBu özet, teknik makalenin son kullanıcı / yönetici seviyesinde anlaşılması "
        "için hazırlanmıştır; resmi bilimsel iddiaların tam metni İngilizce makalededir."
    ).italic = True

    doc.save(str(OUT))
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
