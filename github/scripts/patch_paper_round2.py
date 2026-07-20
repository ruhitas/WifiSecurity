"""Round-2 paper updates from reviewer evidence (Weak Accept conditions)."""
from __future__ import annotations

import json
import shutil
from pathlib import Path

from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "AWID3_Paper_IEEE_Access.docx"
BACKUP = ROOT / "AWID3_Paper_IEEE_Access_pre_round2.docx"
EVIDENCE = json.loads((ROOT / "data" / "reports" / "reviewer_evidence.json").read_text(encoding="utf-8"))
REPO = "https://github.com/ruhitas/WifiSecurity"
ZENODO = "https://zenodo.org/ (DOI to be minted at camera-ready)"


def set_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for r in paragraph.runs[1:]:
            r.text = ""
    else:
        paragraph.add_run(text)


def replace_all(doc: Document, old: str, new: str) -> int:
    n = 0
    for p in doc.paragraphs:
        if old in p.text:
            set_text(p, p.text.replace(old, new))
            n += 1
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if old in p.text:
                        set_text(p, p.text.replace(old, new))
                        n += 1
    return n


def insert_after(doc: Document, needle: str, lines: list[str]) -> bool:
    body = doc.element.body
    for i, p in enumerate(doc.paragraphs):
        if needle in p.text:
            ref = p._element
            for line in reversed(lines):
                np = doc.add_paragraph(line)
                ref.addnext(np._element)
            return True
    return False


def remove_hybrid_from_table4(doc: Document) -> None:
    for table in doc.tables:
        if not table.rows:
            continue
        hdr = " ".join(c.text for c in table.rows[0].cells)
        if "macro-F1" in hdr and "Lat." in hdr:
            for row in list(table.rows):
                if row.cells and "Hybrid" in row.cells[0].text:
                    table._tbl.remove(row._tr)


def remove_paragraphs_containing(doc: Document, sub: str) -> None:
    for p in list(doc.paragraphs):
        if sub in p.text:
            el = p._element
            el.getparent().remove(el)


def add_appendix(doc: Document) -> None:
    feats = EVIDENCE["feature_count_rationale"]["feature_names"]
    lines = [
        "APPENDIX A. LEAKAGE-CONTROLLED FEATURE SET (34 FIELDS)",
        EVIDENCE["feature_count_rationale"]["why_not_other_counts"],
        "Table A1 lists the exact fields retained after the policy in Section IV-A.",
        "TABLE A1. Thirty-four leakage-controlled AWID3 per-frame features.",
    ]
    insert_after(doc, "REFERENCES", lines)

    # feature table after REFERENCES header block - append at end
    doc.add_paragraph("TABLE A1. Thirty-four leakage-controlled AWID3 per-frame features.")
    tbl = doc.add_table(rows=1, cols=2)
    tbl.rows[0].cells[0].text = "#"
    tbl.rows[0].cells[1].text = "Feature name"
    for i, f in enumerate(feats, 1):
        row = tbl.add_row()
        row.cells[0].text = str(i)
        row.cells[1].text = f

    doc.add_paragraph(
        "APPENDIX B. HYBRID CNN+TRANSFORMER (SUPPLEMENTARY)\n"
        "The hybrid model is reported separately because deep training is evaluated on the "
        "stratified 75/25 hold-out (not 5-fold CV) for computational cost. On that split it "
        "reaches 0.929 macro-F1 and 0.974 accuracy — below XGBoost (0.976 macro-F1, 5-fold CV) "
        "and consistent with tabular-data expectations. A full 5-fold CV evaluation is left to "
        "future work; the released train_hybrid.py script reproduces the hold-out numbers."
    )


def main() -> None:
    if not BACKUP.exists():
        shutil.copy2(DOCX, BACKUP)

    ev = EVIDENCE
    naive = ev["leave_group_out_naive"]
    radio = next(r for r in naive if "radio" in r["group"])
    perm = ev["permutation_importance_naive"][0]
    cis = ev["bootstrap_cis"]
    ab15 = next(r for r in ev["ablation_cv_std"] if r["k_features"] == 15)
    ab20 = next(r for r in ev["ablation_cv_std"] if r["k_features"] == 20)

    doc = Document(str(DOCX))

    # --- Major 1: leakage attribution section ---
    leakage_block = [
        "C2) LEAKAGE SOURCE ATTRIBUTION. t-SNE shows cluster separation under naive sampling, "
        "but not which features carry the signal. We therefore ran leave-one-feature-group-out "
        "and permutation-importance analyses on the naive across-capture split (logistic regression, "
        "3-fold CV; n=25,000 stratified sample). Removing the radio group (RSSI, channel, rate) "
        f"drops macro-F1 from {naive[0]['macro_f1']:.4f} to {radio['macro_f1']:.4f} "
        f"(Δ={radio['delta_f1']:.4f}), while the same removal on the curated same-capture split "
        "changes performance in the opposite direction — radio then encodes attack behaviour, not "
        "capture identity. Permutation importance on the naive split ranks "
        f"{perm['feature']} first (mean ΔF1={perm['importance_mean']:.4f}), confirming that "
        "capture-environment fields — not attack semantics alone — drive the inflated scores. "
        "Figure 13 summarises the group ablation.",
        "FIGURE 13. Leave-one-feature-group-out macro-F1 on naive vs curated splits (LR, 3-fold CV).",
    ]
    insert_after(doc, "FIGURE 10. t-SNE of the curated same-capture split", leakage_block)

    # --- Major 2: 34 feature rationale in IV-A ---
    replace_all(
        doc,
        "Removing sequence fields does not materially change the benchmark (Section VI-C).",
        "The resulting count is deterministic, not tuned: 254 AWID3 numeric fields minus "
        "policy-defined leaky columns minus constant columns in the 74,270-row build yields "
        "exactly 34 (Appendix A, Table A1). Removing sequence fields does not materially "
        "change the benchmark (Section VI-C).",
    )

    # --- Major 3: hybrid to appendix ---
    remove_hybrid_from_table4(doc)
    remove_paragraphs_containing(doc, "Hybrid results are from the stratified")

    # --- Major 4: bootstrap F1 CI ---
    replace_all(
        doc,
        "We also bootstrap macro ROC-AUC on held-out predictions (1,000 resamples); the point estimate is "
        "0.9997 with a 95% interval of [0.9996, 0.9997], consistent with the five-fold CV mean of 0.9998 "
        "in Table 4. The reading is plain.",
        f"We bootstrap both macro-F1 and macro ROC-AUC on held-out predictions (1,000 resamples). "
        f"Macro-F1 is {cis['macro_f1_point']:.4f} with 95% CI [{cis['macro_f1_ci95'][0]:.4f}, "
        f"{cis['macro_f1_ci95'][1]:.4f}]; macro ROC-AUC is {cis['roc_auc_point']:.4f} with 95% CI "
        f"[{cis['roc_auc_ci95'][0]:.4f}, {cis['roc_auc_ci95'][1]:.4f}], consistent with the five-fold "
        f"CV means in Table 4. The reading is plain.",
    )

    # --- Major 5: deployment language ---
    replace_all(doc, "companion streaming implementation", "intended deployment pipeline (implemented in the supplementary code package; offline evaluation reported here)")
    replace_all(doc, "The platform is not a sketch; it runs.", "")
    replace_all(doc, "real-time setting rather than in isolation", "streaming-oriented design; this paper reports offline evaluation only")

    # --- Major 6: AWID2 ---
    replace_all(
        doc,
        "cross-wireless-dataset validation (for example AWID2) and live traffic remain future work.",
        "cross-wireless-dataset validation remains the highest-priority extension: AWID2 would "
        "directly test whether capture-environment leakage generalises beyond AWID3, but was "
        "not available for this study; live traffic validation is likewise future work.",
    )

    # --- Major 7: feature engineering ---
    fe_block = [
        "D. FROM RAW FRAMES TO 34 PER-FRAME FEATURES",
        "AWID3 ships 254 Wireshark-derived fields per frame. Our build script (build_awid3.py) "
        "resolves labels by name, applies the leakage policy of Section IV-A, aligns the feature "
        "union across attack folders, fills missing values with zero, and drops columns with zero "
        "variance in the final 74,270-row sample — yielding 34 behavioural fields (Appendix A). "
        "The companion streaming stack aggregates frames into tumbling windows and computes 300+ "
        "window-level statistics for live scoring; the offline benchmark in this paper uses the "
        "per-frame 34-field schema for reproducibility and direct comparison with prior AWID3 work.",
    ]
    insert_after(doc, "Re-running it reproduces the dataset exactly, which is the point.", fe_block)

    # --- Major 8: SHAP locals ---
    replace_all(
        doc,
        "Figure 11 gives the global view: a SHAP beeswarm for the attack-versus-benign decision, where radio, timing and frame-type features carry the weight. Explanation adds operator value without contradicting the numbers.",
        "Figure 11 gives the global beeswarm view. Figures 14–16 add local bar explanations for "
        "Normal, Krack and Deauth windows (true labels), showing class-specific drivers rather "
        "than aggregate patterns alone. Explanation adds operator value without contradicting the numbers.",
    )
    insert_after(
        doc,
        "FIGURE 11. SHAP beeswarm",
        [
            "FIGURE 14. Local SHAP explanation — Normal traffic.",
            "FIGURE 15. Local SHAP explanation — Krack.",
            "FIGURE 16. Local SHAP explanation — Deauth.",
        ],
    )

    # --- Minor fixes ---
    replace_all(doc, "Reported metric", "Reported metric (Acc. or macro-F1 as stated)")
    replace_all(
        doc,
        "TABLE 4. Leakage-controlled 14-class benchmark (5-fold cross-validation, mean ± std), ranked by macro-F1.",
        "TABLE 4. Leakage-controlled 14-class benchmark (5-fold CV, mean ± std), ranked by macro-F1. "
        "Latency: median of 1,000 single-row predict() calls (batch size = 1) after warmup.",
    )
    replace_all(
        doc,
        "Macro-F1 climbs from 0.23 at five features to 0.95 at fifteen, then flattens by twenty at 0.968; the full forty-feature set adds nothing measurable.",
        f"Macro-F1 climbs to {ab15['macro_f1_mean']:.3f}±{ab15['macro_f1_std']:.3f} at fifteen features "
        f"and {ab20['macro_f1_mean']:.3f}±{ab20['macro_f1_std']:.3f} at twenty (5-fold CV); the full "
        f"34-feature set adds nothing measurable (Figure 8, with error bars in supplementary fig_ablation_cv_std.png).",
    )
    replace_all(
        doc,
        "Four transport time-delta",
        "Four transport time-delta",
    )
    replace_all(
        doc,
        "tau <- quantile(s_b, 1 - alpha)          // threshold set on benign only (~5% FPR)",
        "tau <- quantile(s_b, 1 - alpha)          // 95th percentile of benign scores (~5% FPR on X_b)",
    )
    replace_all(
        doc,
        f"Source code: {REPO}\nArchive (Zenodo DOI at camera-ready): {ZENODO}",
        f"Source code: {REPO}\nArchive (Zenodo DOI at camera-ready): {ZENODO}",
    )

    add_appendix(doc)
    doc.save(str(DOCX))
    print(f"Patched -> {DOCX}")


if __name__ == "__main__":
    main()
